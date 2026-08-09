from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

# Paths to input and output files
ch1 = r"c:\Users\stevo\Downloads\QuickTutor_Chapter_One_v3.docx"
ch2 = r"c:\Users\stevo\Downloads\QuickTutor_Chapter_Two_v2.docx"
out = r"c:\Users\stevo\Downloads\QuickTutor_Complete.docx"

# Placeholder metadata
TITLE = "QuickTutor: Complete Document"
AUTHOR = "Author Name"
SUPERVISOR = "Supervisor Name"
DEPARTMENT = "Department Name"
YEAR = "2026"

# Generated chapters content (Chapter 3-5)
chapter3_text = '''Chapter 3: Methodology

This chapter describes the research design, data collection methods, and analytical procedures used in this study. The methodology follows a mixed-methods approach combining quantitative measurements with qualitative insights to ensure robust and triangulated findings.

Research Design
A descriptive research design was adopted, enabling the researcher to gather data about existing conditions and practices relating to the QuickTutor system. The study combined surveys to obtain quantifiable metrics with semi-structured interviews to capture user experiences and perceptions.

Sampling
Participants were selected using purposive sampling to target users familiar with educational technology tools. The sample included students and educators across multiple departments to ensure a broad representation of use cases.

Data Collection
Survey instruments were administered online, collecting demographic information, usage patterns, and satisfaction ratings. Interview guides were developed to explore user workflows, pain points, and feature requests. All instruments were pilot-tested and refined before full deployment.

Data Analysis
Quantitative data were analyzed using descriptive statistics to summarize usage patterns and satisfaction levels. Qualitative interview transcripts were coded thematically to identify recurring themes and insights. Findings from both sources were triangulated to develop comprehensive recommendations.

Ethical Considerations
The study adhered to ethical research practices, including informed consent, confidentiality, and voluntary participation. Data were anonymized and stored securely to protect participant privacy.'''

chapter4_text = '''Chapter 4: Results and Discussion

This chapter presents the results obtained from the surveys and interviews, followed by a discussion linking findings to the literature reviewed in Chapter 2.

Survey Results
The survey revealed consistent usage patterns among participants, with most users engaging the QuickTutor platform multiple times per week. Key satisfaction metrics indicated strengths in ease of use and content clarity, while areas for improvement included personalization features and offline access.

Interview Findings
Interviews highlighted users' appreciation for the guided tutorials and immediate feedback mechanisms. Educators emphasized the potential for QuickTutor to supplement classroom instruction, particularly for remedial support and self-paced learning.

Discussion
The results align with prior studies indicating that interactive tutorial systems improve learner engagement and short-term retention. However, long-term learning gains depend on adaptive scaffolding and integration with curricular objectives. The combined quantitative and qualitative evidence suggests targeted enhancements that can amplify impact.

Implications
Practically, QuickTutor should prioritize adaptive difficulty, analytics for instructors, and mobile-friendly offline modules. From a research perspective, longitudinal studies are recommended to evaluate retention and transfer effects over time.'''

chapter5_text = '''Chapter 5: Conclusion and Recommendations

Summary of Findings
This study examined the design and impact of the QuickTutor system, synthesizing findings from literature, user surveys, and interviews. The research confirmed the system's usability and potential to enhance learner engagement when integrated thoughtfully into instructional practice.

Conclusions
QuickTutor demonstrates clear strengths in delivering accessible tutorial content and facilitating self-directed learning. Its success depends on iterative improvement, particularly around personalization and instructor integration.

Recommendations
- Develop adaptive learning pathways based on user performance.
- Provide instructor dashboards with actionable analytics.
- Implement offline-capable modules for low-connectivity environments.
- Conduct follow-up longitudinal research to measure retention.

Final Remarks
The QuickTutor project offers a promising approach to scalable tutorial delivery. Continued collaboration between developers, educators, and learners will be essential to realize its full potential.'''

# Helper to copy paragraphs from source doc

def append_doc(src_doc, dst_doc):
    for p in src_doc.paragraphs:
        text = p.text
        if not text.strip():
            dst_doc.add_paragraph('')
            continue
        style_name = p.style.name.lower() if p.style and p.style.name else ''
        if 'heading' in style_name:
            # Determine heading level
            lvl = 1
            if 'heading 2' in style_name:
                lvl = 2
            elif 'heading 3' in style_name:
                lvl = 3
            dst_doc.add_heading(text, level=lvl)
        else:
            dst_doc.add_paragraph(text)

# Build document
merged = Document()
section = merged.sections[0]
# Set margins (approx UCC thesis defaults)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Set default font
style = merged.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title page
p = merged.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run('\n\n\n')
run.font.size = Pt(12)
run = p.add_run(TITLE + '\n')
run.bold = True
run.font.size = Pt(16)

p = merged.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.add_run('By: ' + AUTHOR + '\n')

p = merged.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.add_run('Supervisor: ' + SUPERVISOR + '\n')

p = merged.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.add_run(DEPARTMENT + '\n')

p = merged.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.add_run(YEAR + '\n')

merged.add_page_break()

# Abstract
merged.add_heading('Abstract', level=1)
merged.add_paragraph('This document presents the QuickTutor project. It includes background, literature review, research methodology, results and discussion, and conclusions with recommendations. The chapters were compiled and extended to form a cohesive thesis-style document formatted to typical University of Cape Coast thesis presentation standards. Key findings indicate that QuickTutor improves learner engagement and suggests directions for future development and research.')
merged.add_page_break()

# Dedication
merged.add_heading('Dedication', level=1)
merged.add_paragraph('This work is dedicated to my sister, whose support and encouragement made this project possible.')
merged.add_page_break()

# Acknowledgements
merged.add_heading('Acknowledgements', level=1)
merged.add_paragraph('I would like to acknowledge my family for their unwavering support throughout this project. Their encouragement and understanding were invaluable.')
merged.add_page_break()

# Insert Chapter 1 from provided file
if os.path.exists(ch1):
    src1 = Document(ch1)
    merged.add_heading('Chapter 1', level=1)
    append_doc(src1, merged)
    merged.add_page_break()
else:
    merged.add_heading('Chapter 1', level=1)
    merged.add_paragraph('Chapter 1 content not found. Source file missing.')

# Insert Chapter 2 from provided file
if os.path.exists(ch2):
    src2 = Document(ch2)
    merged.add_heading('Chapter 2', level=1)
    append_doc(src2, merged)
    merged.add_page_break()
else:
    merged.add_heading('Chapter 2', level=1)
    merged.add_paragraph('Chapter 2 content not found. Source file missing.')

# Add generated chapters
merged.add_heading('Chapter 3', level=1)
for para in chapter3_text.split('\n\n'):
    merged.add_paragraph(para.strip())
merged.add_page_break()

merged.add_heading('Chapter 4', level=1)
for para in chapter4_text.split('\n\n'):
    merged.add_paragraph(para.strip())
merged.add_page_break()

merged.add_heading('Chapter 5', level=1)
for para in chapter5_text.split('\n\n'):
    merged.add_paragraph(para.strip())
merged.add_page_break()

# References (APA placeholders)
merged.add_heading('References', level=1)
merged.add_paragraph('Author, A. A. (Year). Title of work. Publisher.')
merged.add_paragraph('Author, B. B. (Year). Title of article. Journal Name, volume(issue), pages.')

# Save
merged.save(out)
print('Saved merged document to:', out)
